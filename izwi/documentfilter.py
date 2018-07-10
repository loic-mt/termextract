import logging
import nltk
import codecs
import re
from designpatterns import Visitor 
import languagedetection
import time
import sys
import timeout
import os.path
from nltk.tokenize.treebank import TreebankWordTokenizer

debug = False

word_tokenize = TreebankWordTokenizer().tokenize

class sentence(object):
    def __init__(self,text,offset):
        self.text = text
        self.length = len(text)
        self.offset = offset
        self.char_indices = []
        self.tokens = word_tokenize(text)
        offset = 0
        if text:
            for t in self.tokens:
                self.char_indices.append(offset)
                offset += 1 + len(t)
            self.char_indices.append(offset)
    def __len__(self):
        return self.length
    def token_length(self):
        return len(self.tokens)
    def char_to_token_index(self,char_index):
        return next(ci for ci in self.char_indices if ci == char_index)

    def offset_at_char_range(self, start, end):
        """Convert the given character indices to token indices."""
        if debug:
            print str(self.char_indices)

        if start not in self.char_indices or end not in self.char_indices:
            return None

        start = self.offset + self.char_indices.index(start)
        end = self.offset + self.char_indices.index(end) - 1
        return (start, end)

    def __add__(self,other):
        s = sentence(self.text + " " + other.text,self.offset)
        # we could also avoid re-doing the tokenizing in __init__
        return s
    def tokenized_sentence(self):
        return ' '.join(self.tokens)

class paragraph(object):
    def __init__(self,text,offset):
        self.offset = offset
        self.sentences = []
        self.length = 0
        if text:
            raw_sentences = nltk.sent_tokenize(text)
            for rs in raw_sentences:
                s = sentence(rs, offset)
                offset += s.token_length()
                self.sentences.append(s)
                self.length += len(s)
        self.next_offset = offset
    def __len__(self):
        return self.length
    def __add__(self,other):
        p = paragraph("",self.offset)
        p.sentences = self.sentences
        p.sentences.extend(other.sentences)
        p.length = self.length + other.length
        return p
    def get_sentence_lengths(self):
        sent_lengths = []
        for s in self.sentences:
            sent_lengths.append(len(s))
        return sent_lengths

class izwiDocument:
    def __init__(self,thisfilename):
        self.language = None
        self.filename = thisfilename
        self.paragraphs = []
        self.offset = 0
    def __str__(self):
        return self.language+"\t"+self.filename

    def set_language(self,thislang):
        self.language = thislang
    def add_paragraph(self,p):
        newpar = paragraph(p,self.offset)
        self.paragraphs.append(newpar)
        self.offset = newpar.next_offset
    def get_language(self):
        return self.language
    def get_filename(self):
        return self.filename
    def get_paragraphs(self):
        return self.paragraphs
    def get_paragraph_lengths(self):
        par_lengths = []
        for par in self.paragraphs:
            par_lengths.append(len(par))
        return par_lengths
    def accept(self, visitor):
        visitor.visit(self)
    def nb_paragraphs(self):
        return len(self.paragraphs)
    def get_tokens(self):
        tokens = []
        for p in self.paragraphs:
            for s in p.sentences:
                tokens.extend(s.tokens)
        return tokens

class documentFilter(object):
    def __init__(self):
        pass

    def extract_doc(self,filepath,fileformat):
        if not fileformat:
            _root, ext = os.path.splitext(filepath)
            fileformat = ext[1:]

        if (fileformat == "docx"):
            mydoc = self.extract_doc_from_docx(filepath)
        elif (fileformat == "pdf"):
            mydoc = self.extract_doc_from_pdf(filepath)
        elif (fileformat == "txt"):
            mydoc = self.extract_doc_from_txt(filepath)
        else:
            print "ERROR FILEFORMAT NOT SUPPORTED"
            exit(1)
        return mydoc

    def detect_encoding(self, thisstr):
        from chardet.universaldetector import UniversalDetector
        detector = UniversalDetector()

        for line in thisstr.split('\n'):
            detector.feed(line)
            if detector.done: break
        detector.close()
        return detector.result

    def extract_doc_from_docx(self, thisfile):
        from docx import Document

        myfile = open(thisfile)
        document = Document(myfile)
        myfile.close()

        mydoc = izwiDocument(thisfile)

        content = u""
        for p in document.paragraphs:
            text = p.text.strip()
            if len(text) > 1:
                mydoc.add_paragraph(text)
                if len(content) < languagedetection.MAX_LENGTH:
                    content += text
        lang = languagedetection.most_probable_language(content)
        mydoc.set_language(lang)
        return mydoc

    @timeout.timeout(5)
    def get_content_from_pdf_page(self, rsrcmgr, codec, laparams, page):
        from cStringIO import StringIO
        from pdfminer.converter import TextConverter
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter

        retstr = StringIO()
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        interpreter.process_page(page)
        content = retstr.getvalue().decode(codec)
        retstr.close()
        device.close()

        return content

    def extract_doc_from_pdf(self, thisfile):
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.layout import LAParams
        from pdfminer.pdfpage import PDFPage
        from cStringIO import StringIO

        rsrcmgr = PDFResourceManager()

        codec = 'utf-8'
        laparams = LAParams()

        fp = file(thisfile, 'rb')

        password = ""
        maxpages = 0
        caching = True
        pagenos = set()
        content = u""

        mydoc = izwiDocument(thisfile)
        pagenum = 1

        for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
            try:
                page_content = self.get_content_from_pdf_page(rsrcmgr, codec, laparams, page)
            except:
                logging.warning("Timeout on page %(num)s in %(name)s: skipping",
                                {"num": pagenum, "name": thisfile})
                continue
            page_content = page_content.strip()
            mydoc.add_paragraph(page_content)
            if len(content) < languagedetection.MAX_LENGTH:
                content += page_content
                pagenum += 1
        fp.close()

        lang = languagedetection.most_probable_language(content)
        mydoc.set_language(lang)
        return mydoc

    def extract_doc_from_txt(self, thisfile):
        mydoc = izwiDocument(thisfile)

        with open(thisfile, 'r') as f:
            para = ""
            content = u""
            for line in f.readlines():
                line = line.strip()
                if line:
                    para += line
                elif para:
                    para = para.decode('utf-8')
                    if len(content) < languagedetection.MAX_LENGTH:
                        content += para
                    mydoc.add_paragraph(para)
                    para = ""
                # handle the last line:
            if para:
                para = para.decode('utf-8')
                content += para
                mydoc.add_paragraph(para)
        lang = languagedetection.most_probable_language(content)
        mydoc.set_language(lang)
        return mydoc
